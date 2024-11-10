from fastapi import FastAPI, HTTPException, Query
import openai
from docx import Document
import os
import logging
from dotenv import load_dotenv
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import re
from fastapi.middleware.cors import CORSMiddleware

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # フロントエンドのオリジンに置き換え
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# .envファイルの内容を読み込み
# load_dotenv()→本番環境はHerokuの環境変数を使用

# APIキーを取得
openai.api_key = os.getenv("OPENAI_API_KEY")

# ログ設定
logging.basicConfig(level=logging.INFO)

# 文書検索用関数
def search_relevant_text(query, texts):
    # TF-IDFベクトル化
    vectorizer = TfidfVectorizer().fit_transform([query] + texts)
    vectors = vectorizer.toarray()

    # コサイン類似度計算
    cosine_similarities = cosine_similarity(vectors[0:1], vectors[1:]).flatten()
    top_indices = cosine_similarities.argsort()[-3:][::-1]  # 類似度上位3つを取得
    return [texts[i] for i in top_indices]

# 中央値を表から直接取得する関数
def extract_ev_ebitda_median(document):
    # 1～3桁以上の数値（小数点含む）を抽出するパターン
    pattern = re.compile(r"(\d{1,3}(?:\.\d+)?)倍")
    
    for table_index, table in enumerate(document.tables):
        logging.info(f"Checking table {table_index + 1}")
        for row_index, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            logging.info(f"Row {row_index + 1}: {cells}")  # 各行のセル内容をログに出力

            # "中央値"が見つかったら、右端のセルを確認
            if "中央値" in cells[0]:  # 中央値行の最初のセルをチェック
                logging.info(f"'中央値' found in row {row_index + 1}")
                # 右端のセルがEV/EBITDA倍率
                ev_ebitda_text = cells[-1]
                logging.info(f"EV/EBITDA cell text: {ev_ebitda_text}")
                
                match = pattern.search(ev_ebitda_text)
                if match:
                    logging.info(f"Extracted median EV/EBITDA: {match.group(0)}")
                    return match.group(0)  # 倍率形式（例: 8.4倍）で返す

    logging.warning("企業価値/EBITDA倍率の中央値が見つかりません。")
    return None

@app.get("/summarize")
async def summarize_preset_file(industry: str = Query(None)):
    # `industry`が指定されていない場合にエラーメッセージを返す
    if industry is None:
        raise HTTPException(status_code=400, detail="レポートが表示されていません。業界名を指定してください。")
    
    # industryの値をログ出力して確認
    logging.info(f"Received industry parameter: {industry}")

    try:
        file_path =  f"app/{industry}.docx"  # 選択した業界名に基づいてファイルパスを設定

        # ファイルの存在確認
        if not os.path.exists(file_path):
            logging.error(f"File not found: {file_path}")  # ファイルパスをログに出力
            raise HTTPException(status_code=404, detail="指定されたファイルが見つかりません。")

        # ドキュメント全体の内容を取得し、段落ごとに分割
        document = Document(file_path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

        # 各要素の検索
        queries = {
            "current_situation": "業界の現状を説明してください。",
            "future_outlook": "業界の将来の見立てを説明してください。",
            "investment_advantages": "業界への投資メリットを教えてください。",
            "investment_disadvantages": "業界への投資デメリットを教えてください。",
            "value_up_hypothesis": "業界におけるDXによるバリューアップ仮説を説明してください。",
        }

        # summariesの初期化
        summaries = {
            "current_situation": "",
            "future_outlook": "",
            "investment_advantages": "",
            "investment_disadvantages": "",
            "value_up_hypothesis": ""
        }

        # 各要素に関連する段落を検索して要約
        for key, query in queries.items():
            relevant_paragraphs = search_relevant_text(query, paragraphs)
            relevant_text = "\n".join(relevant_paragraphs)
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {
                        "role": "user",
                        "content": f"{relevant_text}\n\n上記の内容を要約し、以下の質問に回答してください：{query} トークン数は50以内でお願いします。"
                    }
                ]
            )
            summaries[key] = response['choices'][0]['message']['content'].strip()

        # 中央値をWordファイルの表から直接抽出
        ev_ebitda_median = extract_ev_ebitda_median(document)
        if ev_ebitda_median is None:
            raise HTTPException(status_code=404, detail="企業価値/EBITDA倍率の中央値が見つかりません。")

        # 結果を返す
        return {
            "current_situation": summaries["current_situation"],
            "future_outlook": summaries["future_outlook"],
            "investment_advantages": summaries["investment_advantages"],
            "investment_disadvantages": summaries["investment_disadvantages"],
            "value_up_hypothesis": summaries["value_up_hypothesis"],
            "ev_ebitda_median": ev_ebitda_median
        }

    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="指定されたファイルが見つかりません。")
    except openai.error.OpenAIError as e:
        raise HTTPException(status_code=500, detail=f"OpenAI APIエラー: {str(e)}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))